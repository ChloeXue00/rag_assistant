"""
generate_report.py
生成《行研知识库助手 · 技术要点与学习价值分析》Word 文档
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── 颜色常量 ────────────────────────────────────────────────
C_DARK_BLUE  = RGBColor(0x1A, 0x37, 0x6C)   # 深海蓝（标题）
C_MID_BLUE   = RGBColor(0x2E, 0x6B, 0xC4)   # 中蓝（二级标题）
C_ACCENT     = RGBColor(0x00, 0x9A, 0xD9)   # 天蓝（三级标题 / 强调）
C_GRAY       = RGBColor(0x59, 0x59, 0x59)   # 深灰（正文）
C_LIGHT_GRAY = RGBColor(0xF2, 0xF5, 0xFA)   # 浅蓝灰（表格表头底色）
C_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
C_CODE_BG    = RGBColor(0xF5, 0xF5, 0xF5)


def set_cell_bg(cell, hex_color: str):
    """设置表格单元格背景色（需操作底层 XML）"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), kwargs.get("val", "single"))
        tag.set(qn("w:sz"), kwargs.get("sz", "4"))
        tag.set(qn("w:color"), kwargs.get("color", "BFCFE8"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_heading(doc, text, level=1, color=None, size=None, bold=True, space_before=12, space_after=6):
    """统一添加标题段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    if level == 1:
        run.font.color.rgb = color or C_DARK_BLUE
        run.font.size = Pt(size or 18)
    elif level == 2:
        run.font.color.rgb = color or C_MID_BLUE
        run.font.size = Pt(size or 14)
    elif level == 3:
        run.font.color.rgb = color or C_ACCENT
        run.font.size = Pt(size or 12)
    return p


def add_body(doc, text, color=C_GRAY, size=10.5, space_before=2, space_after=3, indent=False):
    """添加正文段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.color.rgb = color
    run.font.size      = Pt(size)
    return p


def add_bullet(doc, text, level=0, color=C_GRAY, size=10.5):
    """添加带符号的条目段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    bullet = "●" if level == 0 else "○"
    indent = Cm(0.5 * (level + 1))
    p.paragraph_format.left_indent = indent
    run = p.add_run(f"{bullet}  {text}")
    run.font.color.rgb = color
    run.font.size      = Pt(size)
    return p


def add_divider(doc, color="2E6BC4"):
    """添加水平分隔线"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def make_table(doc, headers, rows, col_widths=None, header_bg="1A376C"):
    """
    创建带样式的表格。
    headers: 列标题列表
    rows:    数据行列表（每行为列表）
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # 表头行
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = C_WHITE
        run.font.size = Pt(10)

    # 数据行
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = "F2F5FA" if ri % 2 == 0 else "FFFFFF"
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            # 第一列加粗
            run = p.add_run(cell_text)
            run.font.size = Pt(9.5)
            run.font.color.rgb = C_GRAY
            if ci == 0:
                run.bold = True
                run.font.color.rgb = C_MID_BLUE

    # 设置列宽
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table


# ════════════════════════════════════════════════════════════
#  开始构建文档
# ════════════════════════════════════════════════════════════

doc = Document()

# 全局页边距
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(3.0)

# 全局默认字体
style = doc.styles["Normal"]
style.font.name = "微软雅黑"
style.font.size = Pt(10.5)
style.font.color.rgb = C_GRAY

# ── 封面 ────────────────────────────────────────────────────
p_cover = doc.add_paragraph()
p_cover.paragraph_format.space_before = Pt(60)
p_cover.paragraph_format.space_after  = Pt(8)
p_cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_cover.add_run("行研知识库助手")
r.bold = True; r.font.size = Pt(28); r.font.color.rgb = C_DARK_BLUE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(6)
r2 = p2.add_run("技术要点与 AI 应用开发学习价值分析")
r2.font.size = Pt(16); r2.font.color.rgb = C_MID_BLUE

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(60)
r3 = p3.add_run("基于 RAG 架构的行业研究报告智能问答系统")
r3.font.size = Pt(12); r3.font.color.rgb = C_ACCENT

add_divider(doc, "2E6BC4")

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_before = Pt(10)
r_meta = meta.add_run("技术栈：Streamlit · pdfplumber · sentence-transformers · ChromaDB · Claude API")
r_meta.font.size = Pt(10); r_meta.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
#  第一章  项目全景与技术架构
# ════════════════════════════════════════════════════════════
add_heading(doc, "一、项目全景与技术架构", level=1)
add_body(doc,
    "本项目是一套完整的 RAG（Retrieval-Augmented Generation，检索增强生成）应用。"
    "用户上传行业研究报告 PDF，系统自动建立语义索引；提问时先检索最相关段落，再将其作为"
    "上下文发给 Claude，最终输出有据可查的专业回答。整个流程覆盖了现代 AI 应用开发的完整链路。"
)

add_heading(doc, "1.1  RAG 完整数据流", level=2)
flow_steps = [
    ("① 文档输入", "用户上传 PDF → pdfplumber 按页提取文本"),
    ("② 文本清洗", "去除多余空白、合并断行、压缩空行"),
    ("③ 智能切块", "按自然段落优先切分，超长段硬切；相邻块 50 字重叠"),
    ("④ 向量化",   "sentence-transformers 将每块文本编码为 384 维浮点向量"),
    ("⑤ 持久化",   "ChromaDB 以余弦空间存储向量 + 文本 + metadata"),
    ("⑥ 检索",     "用户问题同样向量化，HNSW 索引找 Top-5 相似段落"),
    ("⑦ 生成",     "检索结果 + 问题 + 对话历史组装成 prompt，发送 Claude API"),
    ("⑧ 溯源展示", "回答附带来源文件名、段落编号、相关度百分比"),
]
make_table(doc,
    headers=["步骤", "说明"],
    rows=flow_steps,
    col_widths=[3.5, 12.5],
)

doc.add_paragraph()
add_heading(doc, "1.2  项目文件职责划分", level=2)
file_roles = [
    ("app.py",              "Streamlit 主程序，负责 UI 渲染、session 状态管理、流程编排"),
    ("utils/pdf_parser.py", "PDF 解析层：文本提取 → 清洗 → 切块，纯 Python 无副作用"),
    ("utils/embedder.py",   "向量层：模型推理、ChromaDB 增删查、来源统计"),
    ("utils/chat.py",       "LLM 层：system prompt、RAG context 构建、Claude API 调用、历史管理"),
    ("requirements.txt",    "依赖声明，锁定主要库最低版本"),
    ("chroma_db/（运行时）", "ChromaDB 自动创建的本地持久化目录"),
]
make_table(doc,
    headers=["文件 / 目录", "职责"],
    rows=file_roles,
    col_widths=[5.0, 11.0],
)

doc.add_paragraph()
add_divider(doc)

# ════════════════════════════════════════════════════════════
#  第二章  六大核心技术模块详解
# ════════════════════════════════════════════════════════════
add_heading(doc, "二、六大核心技术模块详解", level=1)

# ── 2.1 文档解析 ─────────────────────────────────────────
add_heading(doc, "2.1  文档解析与文本切块（pdf_parser.py）", level=2)
add_body(doc, "这是 RAG 质量的基础。解析质量直接决定后续检索的天花板。")

add_heading(doc, "技术要点", level=3)
parse_points = [
    ("pdfplumber 解析",   "逐页调用 page.extract_text()，自动处理多栏、嵌入字体；在每页文本前插入 [第N页] 标记，为后续溯源埋下锚点"),
    ("文本清洗",          "统一换行符、压缩连续空行、去除行首尾空格；保留自然段落结构（双换行）供后续切割"),
    ("优先段落边界切割",  "用 re.split(r'\\n\\n+', text) 按自然段分割，最大程度保留语义完整性，避免在句子中间截断"),
    ("硬切兜底",          "单段落超过 chunk_size 时调用 _hard_split()，确保每块不超上限"),
    ("Overlap 重叠",      "相邻块保留 50 字重叠（可配置），避免关键信息恰好落在块边界被截断，提升跨块问题的检索召回"),
    ("Metadata 注入",     "每块携带 source（文件名）和 chunk_id（序号），在向量库中作为 metadata 存储，支持精确溯源"),
]
make_table(doc,
    headers=["知识点", "具体实现与意义"],
    rows=parse_points,
    col_widths=[4.0, 12.0],
)

doc.add_paragraph()

# ── 2.2 文本向量化 ───────────────────────────────────────
add_heading(doc, "2.2  文本向量化与语义检索（embedder.py）", level=2)
add_body(doc,
    "向量化是 RAG 的核心引擎。将文字转换为高维数字向量，使「语义相近的文本距离相近」，"
    "实现传统关键词搜索无法做到的语义理解。"
)

add_heading(doc, "技术要点", level=3)
embed_points = [
    ("Embedding 模型选型",   "paraphrase-multilingual-MiniLM-L12-v2：支持50+语言（含中文），384维向量，模型轻量（约400MB）；推理速度快，适合本地部署"),
    ("批量推理优化",          "model.encode(texts) 批量处理整个文件的所有块，比逐条推理快3-5倍；show_progress_bar=False 避免日志污染 Streamlit"),
    ("@st.cache_resource",   "模型加载用 Streamlit 资源缓存装饰器，保证多次 rerun 只加载一次，节省数十秒启动时间"),
    ("ChromaDB 余弦相似度",   "创建集合时指定 hnsw:space=cosine，余弦距离对文本语义匹配比欧氏距离更稳定（不受向量模长影响）"),
    ("HNSW 近似最近邻索引",   "ChromaDB 内置 HNSW 算法，百万级向量检索毫秒级响应；trade-off：牺牲极少精度换取指数级速度提升"),
    ("Upsert 去重策略",       "以 sha256(文件名)[:8]+chunk_id 作为唯一 ID，重复上传同一文件只更新而非重复写入，保证数据一致性"),
    ("按来源过滤检索",        "query() 支持 where={source: {$in: [...]}} 过滤，可实现「只在某几份报告中搜索」的精确检索"),
    ("持久化存储",            "PersistentClient 将数据写入磁盘 ./chroma_db，进程重启后数据完整保留"),
]
make_table(doc,
    headers=["知识点", "具体实现与意义"],
    rows=embed_points,
    col_widths=[4.5, 11.5],
)

doc.add_paragraph()

# ── 2.3 Prompt 工程 ──────────────────────────────────────
add_heading(doc, "2.3  Prompt 工程与 Claude API 调用（chat.py）", level=2)
add_body(doc,
    "Prompt 设计决定 LLM 输出质量。本项目采用结构化 prompt 模板，"
    "将系统角色、检索上下文、对话历史、用户问题分层组织。"
)

add_heading(doc, "技术要点", level=3)
prompt_points = [
    ("System Prompt 设计",     "明确角色（行研助手）+ 行为规范（优先用资料、标注来源、格式化输出），约束模型不随意臆造"),
    ("RAG Context 格式化",     "检索结果格式化为「[资料N] 来源：《文件名》第X段 （相关度Y%）\\n正文」，编号便于模型引用，相关度帮助模型判断可信度"),
    ("System/User 消息分离",   "system 字段传角色设定，user 消息传上下文+问题；清晰的角色分离是 Claude API 最佳实践"),
    ("多轮对话历史管理",        "conversation_history 列表按 {role, content} 格式累积，每轮追加后传给 API，实现跨轮记忆"),
    ("历史瘦身策略",            "历史中只存用户原始问题（不存检索上下文），避免每轮历史包含大量重复文档内容导致 token 爆炸"),
    ("环境变量读取 API Key",    "os.environ.get(ANTHROPIC_API_KEY) 而非硬编码，遵循「12-Factor App」安全实践"),
    ("来源标注展示",            "format_sources_for_display() 将检索块格式化为 markdown，在 Streamlit expander 中折叠展示，不影响主回答阅读"),
]
make_table(doc,
    headers=["知识点", "具体实现与意义"],
    rows=prompt_points,
    col_widths=[4.5, 11.5],
)

doc.add_paragraph()

# ── 2.4 Streamlit 前端 ───────────────────────────────────
add_heading(doc, "2.4  Streamlit 全栈 UI 开发（app.py）", level=2)
add_body(doc,
    "Streamlit 是当前 AI 应用原型开发最主流的 Python 框架，"
    "「每次交互重新执行全部脚本」的模型需要特殊的状态管理思维。"
)

add_heading(doc, "技术要点", level=3)
st_points = [
    ("Session State",          "st.session_state 存储对话历史、已上传文件集合，保证 Streamlit 每次 rerun 后状态不丢失"),
    ("st.cache_resource",      "缓存 SentenceTransformer 模型实例（重量级资源），整个 session 只初始化一次"),
    ("st.chat_message()",      "原生聊天气泡 UI，role=user/assistant 自动配色和图标，无需手写 CSS"),
    ("st.chat_input()",        "底部固定输入框，disabled 参数在知识库为空时自动禁用，引导用户先上传文件"),
    ("st.file_uploader()",     "accept_multiple_files=True 支持批量拖拽上传，type=[pdf] 前端过滤非 PDF 文件"),
    ("st.progress()",          "文件处理进度条，提升多文件上传时的用户体验，避免长时间无反馈"),
    ("st.expander()",          "折叠展示参考来源，默认收起不打扰阅读，需要时展开查看原文"),
    ("st.spinner()",           "检索和 API 调用期间显示加载动画，明确告知用户系统正在工作"),
    ("st.rerun()",             "删除文件后立即刷新页面，让文件列表实时更新；是 Streamlit 触发重渲染的标准方式"),
    ("列布局 st.columns()",    "侧边栏文件列表用 [4,1] 比例分栏，左侧显示文件名，右侧放删除按钮"),
]
make_table(doc,
    headers=["API / 模式", "具体应用与作用"],
    rows=st_points,
    col_widths=[4.5, 11.5],
)

doc.add_paragraph()

# ── 2.5 向量数据库 ───────────────────────────────────────
add_heading(doc, "2.5  向量数据库原理（ChromaDB）", level=2)
add_body(doc,
    "向量数据库是 RAG 架构的存储核心，与传统关系型数据库有本质不同：它存储的是「浮点数组」而非结构化数据，"
    "查询的是「语义相似度」而非精确匹配。"
)

add_heading(doc, "核心概念对照", level=3)
chroma_concepts = [
    ("Collection",  '关系数据库的「表」',  "存储同类向量数据的容器，可按业务划分多个 Collection"),
    ("Document",    '数据库的「行内容」',  "原始文本字符串，与向量一一对应存储"),
    ("Embedding",   '数据库的「索引列」',  "document 对应的浮点数向量，用于相似度计算"),
    ("Metadata",    '数据库的「附加列」',  "任意 key-value 字典，可用于过滤（如 where source=xxx）"),
    ("ID",          '数据库的「主键」',    "唯一标识每条记录，upsert 时用于去重更新"),
    ("HNSW Index",  '数据库的「B+树索引」', "分层可导航小世界图，支持近似最近邻快速搜索"),
]
make_table(doc,
    headers=["ChromaDB 概念", "类比", "说明"],
    rows=chroma_concepts,
    col_widths=[3.5, 4.0, 8.5],
)

doc.add_paragraph()

# ── 2.6 工程化实践 ───────────────────────────────────────
add_heading(doc, "2.6  工程化设计实践", level=2)
add_body(doc, "代码工程质量决定可维护性。本项目遵循了多项主流工程实践：")

eng_points = [
    ("配置常量集中管理",     "CHUNK_SIZE / OVERLAP / TOP_K / MODEL 等全部放在 app.py 顶部，改一处全局生效，避免散落各处的魔法数字"),
    ("分层架构",             "UI层（app.py）/ 解析层（pdf_parser）/ 向量层（embedder）/ LLM层（chat）职责分离，互不耦合"),
    ("临时文件安全处理",     "上传文件写入 tempfile.NamedTemporaryFile，finally 块确保 Path.unlink() 清理，无论成功失败都不留垃圾文件"),
    ("类型注解",             "函数签名使用 List[Dict] / Tuple / Optional 等类型标注，提升 IDE 智能提示和代码可读性"),
    ("错误处理分级",         "ValueError 处理配置错误（API Key 缺失），Exception 兜底处理网络/解析异常，错误信息在 UI 清晰展示"),
    ("函数职责单一",         "每个函数只做一件事（如 _hard_split 专门处理超长切割），便于单元测试和复用"),
    ("requirements.txt",    "声明最低版本而非精确版本（>=），在稳定性与兼容性之间取得平衡"),
]
make_table(doc,
    headers=["实践", "说明"],
    rows=eng_points,
    col_widths=[4.5, 11.5],
)

doc.add_paragraph()
add_divider(doc)

# ════════════════════════════════════════════════════════════
#  第三章  学习价值评估
# ════════════════════════════════════════════════════════════
add_heading(doc, "三、学习价值评估：能学到哪些主流技能？", level=1)
add_body(doc,
    "下表对当前 AI 应用岗位（AI 工程师、后端开发、产品技术）的核心技能要求进行对照，"
    "评估本项目的覆盖深度（★ 入门  ★★ 熟练  ★★★ 深入）。"
)

skill_map = [
    ("RAG 架构",          "★★★",  "完整实现「解析→切块→向量化→检索→生成→溯源」全链路，是理解 RAG 最好的入手项目"),
    ("Prompt Engineering","★★★",  "System prompt 设计、RAG context 格式化、多轮历史管理，三个核心 prompt 技巧均有实践"),
    ("向量数据库",         "★★★",  "ChromaDB 增删查、余弦空间、HNSW索引、metadata过滤，覆盖生产环境常用操作"),
    ("LLM API 调用",      "★★★",  "Claude messages API、流式 vs 阻塞模式、token 控制、环境变量安全管理"),
    ("Embedding 模型",    "★★",   "掌握 sentence-transformers 使用，理解向量维度/批量推理；未涉及模型训练/微调"),
    ("Streamlit 开发",    "★★★",  "session_state / cache_resource / chat组件 / 文件上传 / 布局，涵盖80%日常用法"),
    ("Python 工程实践",   "★★",   "分层架构、类型注解、临时文件处理、错误分级；未涉及测试/CI/CD"),
    ("文档处理",           "★★",   "pdfplumber 解析、文本清洗 regex、切块算法；未涉及图片/表格提取"),
    ("本地部署",           "★★",   "理解环境变量、虚拟环境、依赖管理；未涉及 Docker/K8s 容器化"),
    ("Agent / 工具调用",  "★",    "未涉及，是进阶方向（可在本项目基础上扩展）"),
    ("流式输出",           "★",    "本项目为阻塞式调用；流式（Streaming）是常见进阶需求，需改造 chat.py"),
    ("多模态处理",         "—",    "未涉及（PDF 中的图表目前被忽略），是明确的扩展方向"),
]
make_table(doc,
    headers=["技能领域", "掌握深度", "说明"],
    rows=skill_map,
    col_widths=[3.8, 2.2, 10.0],
)

doc.add_paragraph()

add_heading(doc, "3.1  本项目学习路径建议", level=2)
learning_path = [
    ("第1步：跑通",       "按 README 安装依赖，上传一份 PDF，成功提问并看到来源标注"),
    ("第2步：读懂",       "从 pdf_parser.py → embedder.py → chat.py → app.py 逐文件阅读，理解每个函数的输入输出"),
    ("第3步：改参数",     "修改 CHUNK_SIZE / TOP_K，观察回答质量变化，建立直觉"),
    ("第4步：加功能",     "尝试：① 加流式输出  ② 加多 Collection 支持  ③ 加 PDF 表格提取  ④ 用 OpenAI / Gemini 替换 Claude"),
    ("第5步：替换组件",   "把 ChromaDB 换成 Qdrant/Weaviate，把 sentence-transformers 换成 OpenAI Embedding API，理解接口抽象"),
    ("第6步：生产化",     "加 Docker 部署、加用户认证、加对话持久化到 SQLite，体验从原型到生产的差距"),
]
make_table(doc,
    headers=["阶段", "目标"],
    rows=learning_path,
    col_widths=[3.0, 13.0],
)

doc.add_paragraph()
add_divider(doc)

# ════════════════════════════════════════════════════════════
#  第四章  技术选型对比
# ════════════════════════════════════════════════════════════
add_heading(doc, "四、技术选型对比：为什么选这些库？", level=1)

comparison = [
    ("PDF 解析",    "pdfplumber",      "PyPDF2 / pymupdf",   "pdfplumber 对中文/多栏排版支持更好，可直接提取坐标信息；pymupdf 更快但需 C 编译"),
    ("Embedding",  "sentence-transformers", "OpenAI text-embedding-3", "前者完全本地运行零成本，支持中文；后者效果更强但需付费且有网络延迟"),
    ("向量数据库",  "ChromaDB",        "Qdrant / Weaviate / FAISS", "ChromaDB 最易上手，纯 Python 安装；FAISS 无持久化；Qdrant/Weaviate 功能更强但需单独部署"),
    ("前端框架",    "Streamlit",       "Gradio / FastAPI+Vue", "Streamlit 对 AI 原型最友好，内置 chat 组件；Gradio 更简单但定制性差；FastAPI+Vue 最灵活但开发成本高"),
    ("LLM",        "Claude Sonnet",   "GPT-4o / Gemini Pro", "Claude 中文理解和长文本能力强；三者 API 接口相似，可互换"),
]
make_table(doc,
    headers=["层", "本项目选择", "主流替代方案", "选型理由"],
    rows=comparison,
    col_widths=[2.5, 3.5, 4.5, 5.5],
)

doc.add_paragraph()
add_divider(doc)

# ════════════════════════════════════════════════════════════
#  第五章  已知局限与进阶方向
# ════════════════════════════════════════════════════════════
add_heading(doc, "五、已知局限与进阶扩展方向", level=1)

add_heading(doc, "5.1  当前局限", level=2)
limits = [
    "PDF 中的图表、图片被完全忽略，仅处理纯文本",
    "切块策略较简单，学术论文中的公式、代码块可能被错误切割",
    "对话历史无上限，极长对话后 token 数会超出模型限制（需加截断策略）",
    "没有重排序（Reranking）步骤，检索精度依赖 embedding 模型质量",
    "单用户设计，多用户并发会共享同一个向量库，无隔离",
    "Streamlit 刷新页面即丢失对话（未持久化到数据库）",
]
for l in limits:
    add_bullet(doc, l)

add_heading(doc, "5.2  进阶扩展方向", level=2)
advances = [
    ("Reranking 重排序",  "检索 Top-20 后用 CrossEncoder 精排到 Top-5，显著提升精度（BGE-reranker 等模型）"),
    ("混合检索",          "结合 BM25 关键词检索 + 向量检索，取两者分数加权合并（RRF 算法），弥补纯语义检索的短板"),
    ("图表理解",          "接入多模态模型（Claude Vision / GPT-4V）解读 PDF 中的图表，扩充知识库内容"),
    ("流式输出",          "改用 client.messages.stream() 实现打字机效果，提升大段回答的用户体验"),
    ("Agent 化",          "给模型配备工具调用（搜索、计算、数据库查询），从问答升级为自主任务执行"),
    ("对话持久化",        "将 session_state 改为写入 SQLite / Redis，支持历史对话回溯和多用户隔离"),
    ("评估体系",          "引入 RAGAS 框架评估「检索召回率」「答案忠实度」「答案相关性」三个 RAG 核心指标"),
]
make_table(doc,
    headers=["方向", "说明"],
    rows=advances,
    col_widths=[4.0, 12.0],
)

doc.add_paragraph()
add_divider(doc)

# ════════════════════════════════════════════════════════════
#  结语
# ════════════════════════════════════════════════════════════
add_heading(doc, "六、总结", level=1)
add_body(doc,
    "这个项目不大，但「麻雀虽小，五脏俱全」。它完整覆盖了当前 AI 应用开发最主流的技术范式——RAG，"
    "并且每一个技术决策都有真实的工程考量在背后。"
)
add_body(doc,
    "对于 AI 应用开发者而言，本项目的学习价值在于：不只是「会用 API」，"
    "而是理解从原始文档到智能问答的每一个环节——文本如何变成向量、向量如何被检索、"
    "检索结果如何转化成高质量回答、状态如何在 Web 应用中跨请求保持。"
    "这些都是构建任何生产级 AI 应用的基础能力。"
)
add_body(doc,
    "建议以本项目为起点，沿「增加 Reranking → 引入混合检索 → Agent 化」的路径持续深化，"
    "逐步掌握从原型到生产级 RAG 系统的完整技能栈。",
    color=C_MID_BLUE,
)

# ── 保存文档 ─────────────────────────────────────────────
output_path = r"C:\Users\xuefx\Documents\rag_assistant\RAG_技术要点与学习价值分析.docx"
doc.save(output_path)
print(f"文档已生成：{output_path}")
